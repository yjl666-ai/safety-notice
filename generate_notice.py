#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
安全隐患整改通知书生成器(N 条隐患 → 1 段描述 + N 条针对性整改)

用法:
  python generate_notice.py <notice.yaml> [output.docx]

字段见同目录的 notice.yaml.example。改 YAML 字段即可改通知内容。

YAML 关键字段:
  - title, to_unit, from_unit, issue_date, context
  - hazards_input: list[str]    原样保留的隐患列表(供审计/留档)
  - hazard_paragraph: str       AI 生成的"一段话"主体(已含日期、单位、违规条款等)
  - requirements: list[str]     N 条针对性整改(与输入隐患一一对应)
  - location, regulation_name, regulation_code, violation_clause
  - deadline_days, penalty_doc, penalty
  - doc_number (可选)
  - photos: list[str]           4 个占位 ["图1","图2","图3","图4"],实际贴图时替换
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cn_font(run, name="宋体"):
    """给 run 设置中文字体。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_title(doc, text, size=22, font="黑体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = True
    set_cn_font(run, font)


def add_centered(doc, text, size=12, font="宋体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_cn_font(run, font)


def add_body(doc, text, indent=False, bold=False, size=12, font="宋体"):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    set_cn_font(run, font)
    return p


def add_right(doc, text, size=12, font="宋体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_cn_font(run, font)


def add_photo_grid(doc, photos: list):
    """渲染 2x2 照片网格占位(实际使用时人工替换为真实图片)。

    表格方式: 2 列 × 2 行,每格内写 "图N" 占位文字,方便 Word 里直接拖图替换。
    """
    if not photos:
        photos = ["图1", "图2", "图3", "图4"]
    # 补齐到 4
    while len(photos) < 4:
        photos.append(f"图{len(photos)+1}")

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置表格宽度
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(7.5)

    cells = [table.rows[0].cells[0], table.rows[0].cells[1],
             table.rows[1].cells[0], table.rows[1].cells[1]]
    for cell, label in zip(cells, photos[:4]):
        cell.text = ""  # 清空默认段落
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x9B, 0x9B, 0x9B)  # 灰色占位
        set_cn_font(run)
        # 加单元格边框
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "BFBFBF")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    doc.add_paragraph()  # 表格后空行


def init_doc():
    doc = Document()
    section = doc.sections[0]
    for attr, val in [
        ("left_margin", Cm(2.5)),
        ("right_margin", Cm(2.5)),
        ("top_margin", Cm(2.5)),
        ("bottom_margin", Cm(2.5)),
    ]:
        setattr(section, attr, val)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)
    return doc


def generate(data: dict, out_path: Path):
    doc = init_doc()

    # 标题
    add_title(doc, data.get("title", "安全隐患整改通知书"))

    # 文号(可选)
    if data.get("doc_number"):
        add_centered(doc, data["doc_number"], size=12)

    # 2x2 照片网格
    add_photo_grid(doc, data.get("photos", []))

    doc.add_paragraph()  # 空行

    # 抬头
    add_body(doc, f"致:{data.get('to_unit', 'XX项目部')}:", indent=False)

    # 主体:一段话(AI 生成的 hazard_paragraph)
    paragraph = data.get("hazard_paragraph", "")
    if not paragraph:
        # 兜底:拼一句
        paragraph = (
            f"{data.get('issue_date', '')},"
            f"{data.get('from_unit', '安监部')}于"
            f"{data.get('context', '日常巡检')}中,"
            f"发现你单位存在安全隐患,违反"
            f"《{data.get('regulation_name', '')}》"
            f"({data.get('regulation_code', '')})"
            f"{data.get('violation_clause', '')},未及时整改。"
        )
    # 多句就用一个段落(不再切段),保持一段到底
    add_body(doc, paragraph, indent=True)

    add_body(doc, "")
    add_body(doc, "现对你单位要求如下:", indent=False)

    # 整改要求:N 条针对性(一一对应)
    for i, req in enumerate(data.get("requirements", []), 1):
        add_body(doc, f"{i}.{req}", indent=False)

    add_body(doc, "")

    # 闭环段
    deadline = data.get("deadline_days", 3)
    from_unit = data.get("from_unit", "安监部")
    review = data.get("review_clause") or f"{from_unit}将于整改完成后组织复查"
    feed_back = data.get("feedback_clause") or f"并将整改结果书面报{from_unit}"
    tail = (
        f"本通知下发之日起 {deadline} 日内完成全部整改,"
        f"{feed_back}。"
        f"{review}。"
        f"逾期将按照《{data.get('penalty_doc', '施工现场奖惩实施细则')}》"
        f"对你单位进行{data.get('penalty', '违约扣款')}处理。"
    )
    add_body(doc, tail, indent=True)

    add_body(doc, "")
    add_body(doc, "")

    # 落款
    add_right(doc, f"发文单位:{from_unit}(盖章)")
    add_right(doc, data.get("issue_date", ""))

    add_body(doc, "")
    add_body(doc, "签收人:_____________  签收日期:_____________")

    doc.save(out_path)


def dump_for_preview(out_path: Path):
    """把生成的 docx 段落提取出来,用于终端预览。"""
    d = Document(out_path)
    lines = []
    for p in d.paragraphs:
        t = p.text
        if t.strip() == "":
            continue
        if len(t) > 60 and not t.lstrip().startswith(("致:", "1.", "2.", "3.", "4.")):
            lines.append(f"　　{t}")
        elif t.lstrip().startswith(("1.", "2.", "3.", "4.")):
            lines.append(t)
        elif "签收" in t or "盖章" in t:
            lines.append("　　　　　　　　　　　　　　　　　　" + t)
        else:
            lines.append(t)
    # 表格(照片网格)
    for table in d.tables:
        for row in table.rows:
            row_text = "    ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                lines.append("┌─" + "─" * 30 + "─┐")
                lines.append("│ " + row_text + " │")
                lines.append("└─" + "─" * 30 + "─┘")
    return "\n".join(lines)


if __name__ == "__main__":
    yaml_path = Path(sys.argv[1] if len(sys.argv) > 1 else "notice.yaml")
    out_path = Path(sys.argv[2] if len(sys.argv) > 2 else "安全隐患整改通知书.docx")

    import yaml  # noqa: 这里才 import 避免顶层依赖
    with open(yaml_path, encoding="utf-8") as f:
        data = yaml.safe_load(f)

    generate(data, out_path)
    print(f"✓ 已生成: {out_path.resolve()}")
    print("---文件预览(终端版)---")
    print(dump_for_preview(out_path))